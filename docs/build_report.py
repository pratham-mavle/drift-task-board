#!/usr/bin/env python3
"""Build the Drift delivery report as a DOCX for deterministic PDF export."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1C2130"
MUTED = "626A7B"
ACCENT = "5B5BD6"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "2D8C73"
AMBER = "9A681F"
RED = "9B1C1C"
BORDER = "D7DBE2"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    """Apply fixed Word table geometry: tblW, tblInd, tblGrid, and every tcW."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    for tag_name, value in (("tblW", total), ("tblInd", indent_dxa)):
        node = tbl_pr.find(qn(f"w:{tag_name}"))
        if node is None:
            node = OxmlElement(f"w:{tag_name}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    table._tbl.replace(old_grid, new_grid)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph, field_name: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=8, color=MUTED)


def set_header_footer(section, left_text: str, right_text: str):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(left_text.upper())
    set_run_font(run, size=8, color=MUTED, bold=True)
    run = p.add_run(" " + right_text)
    set_run_font(run, size=8, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("DRIFT  |  PAGE ")
    set_run_font(run, size=8, color=MUTED, bold=True)
    add_page_field(fp, "PAGE")
    run = fp.add_run(" OF ")
    set_run_font(run, size=8, color=MUTED)
    add_page_field(fp, "NUMPAGES")


def add_paragraph_shading(paragraph, fill: str, border: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str, color=ACCENT):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc, abstract_id: int, num_id: int, fmt: str, text: str):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    lvl.extend((start, num_fmt, lvl_text, lvl_jc, p_pr))
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def add_list_item(doc, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_el))
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_header_footer(section, "Project delivery report", "Pratham Mavle  ·  22 July 2026")

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    sql_style = doc.styles.add_style("SQL Code", 1)
    sql_style.font.name = "Consolas"
    sql_style.font.size = Pt(7)
    sql_style.font.color.rgb = rgb("24272F")
    sql_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    sql_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    sql_style.paragraph_format.left_indent = Pt(4)
    sql_style.paragraph_format.right_indent = Pt(4)
    sql_style.paragraph_format.space_before = Pt(0)
    sql_style.paragraph_format.space_after = Pt(0)
    sql_style.paragraph_format.line_spacing = Pt(8.2)

    create_numbering(doc, 77, 77, "bullet", "•")
    create_numbering(doc, 78, 78, "decimal", "%1.")
    create_numbering(doc, 79, 79, "decimal", "%1.")


def add_cover(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(54)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("PROJECT DELIVERY REPORT")
    set_run_font(run, size=9, color=ACCENT, bold=True)
    run.font.all_caps = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Drift")
    set_run_font(run, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("A calmer Kanban workspace for focused teams")
    set_run_font(run, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("React + TypeScript + Supabase  |  Guest auth + RLS  |  July 2026")
    set_run_font(run, size=9.5, color=MUTED, italic=True)

    image_path = PROJECT_ROOT / "public" / "og.png"
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(17)
        inline = p.add_run().add_picture(str(image_path), width=Inches(6.25))
        inline._inline.docPr.set("descr", "Drift Kanban board product preview")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Prepared by Pratham Mavle")
    set_run_font(run, size=10, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("22 July 2026")
    set_run_font(run, size=9, color=MUTED)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, tone=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    add_paragraph_shading(p, CALLOUT, tone)
    run = p.add_run(label.upper() + "\n")
    set_run_font(run, size=8, color=tone, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_fact_table(doc: Document):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    repeat_header_row(table.rows[0])
    rows = [
        ("Live frontend", "https://drift-task-board.psmavle02.chatgpt.site"),
        ("Public repository", "https://github.com/pratham-mavle/drift-task-board"),
        ("Core stack", "React 19, TypeScript, Vinext/Vite, Supabase, dnd-kit, Lucide"),
        ("Delivery status", "Public Supabase deployment active; anonymous Auth, RLS persistence, and realtime configured"),
    ]
    for idx, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
        set_cell_text(table.cell(idx, 0), label, bold=True, color=DARK_BLUE)
        if idx < 2:
            p = set_cell_text(table.cell(idx, 1), "")
            link_text = "Open live app" if idx == 0 else "View public repository"
            add_hyperlink(p, link_text, value)
        else:
            set_cell_text(table.cell(idx, 1), value)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_feature_table(doc: Document):
    features = [
        ("Kanban workflow", "Four fixed columns, cross-column drag, same-column ordering, keyboard/touch sensors, optimistic rollback", "Built"),
        ("Task management", "Create, edit, delete, descriptions, priority, due dates, status, ordering", "Built"),
        ("Guest accounts", "Automatic Supabase anonymous sign-in with persisted browser session", "Built"),
        ("Team + assignees", "User-owned lightweight member profiles and many-to-many task assignments", "Advanced"),
        ("Comments", "Chronological task comments with timestamps and realtime refresh", "Advanced"),
        ("Activity log", "Append-only trigger-generated history for moves, edits, assignments, labels, and comments", "Advanced"),
        ("Labels + filters", "Reusable labels; search plus priority, assignee, and label filters", "Advanced"),
        ("Due-date signals", "Overdue, due-soon, later, and completed treatments directly on cards", "Advanced"),
        ("Summary + states", "Total/completed/overdue stats; loading, empty, no-results, reconnecting, and error states", "Advanced"),
        ("Responsive UI", "Collapsed navigation, horizontal mobile board, touch drag, full-screen detail sheet", "Built"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 5960, 1200])
    set_table_borders(table)
    for idx, title in enumerate(("Capability", "Implementation", "Status")):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        set_cell_text(table.cell(0, idx), title, bold=True, color=DARK_BLUE, size=9)
    repeat_header_row(table.rows[0])
    for feature, detail, status in features:
        cells = table.add_row().cells
        set_cell_text(cells[0], feature, bold=True, size=9)
        set_cell_text(cells[1], detail, size=8.8)
        p = set_cell_text(cells[2], status, bold=True, color=GREEN if status == "Built" else ACCENT, size=8.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [2200, 5960, 1200])
    doc.add_paragraph()


def add_code_block(doc: Document, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.right_indent = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        add_paragraph_shading(p, "F7F7F9")
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=8.3, color="313641")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_main_content(doc: Document):
    doc.add_heading("Solution at a glance", level=1)
    add_callout(
        doc,
        "Outcome",
        "Drift delivers a polished, team-ready Kanban experience backed by a live Supabase Free Tier project. The public deployment creates an anonymous guest session automatically, persists work in Postgres, enforces per-user access through RLS, and refreshes owned workspace data through Realtime.",
    )
    add_fact_table(doc)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Drift is a single-page task workspace designed to feel closer to Linear, Asana, and Notion than a generic todo list. The experience centers the board, keeps secondary controls quiet, and lets users move from scanning to editing without losing context. A deep charcoal rail anchors the interface; warm neutral surfaces, restrained indigo accents, compact cards, and semantic status colors create hierarchy without visual noise."
    )
    doc.add_paragraph(
        "The frontend calls Supabase directly with its public publishable browser key. The production deployment first restores or creates an anonymous Auth session, then loads only rows allowed by Row Level Security. An explicit browser-local demo path remains available for development, but a failing live connection never silently falls back to local data."
    )

    doc.add_heading("Architecture and data flow", level=2)
    architecture = [
        "Guest identity. Supabase Auth restores the browser session or calls signInAnonymously() once on first launch.",
        "Protected data. The React client reads tasks and collaboration records through the public anon key; RLS evaluates auth.uid() for every request.",
        "Optimistic interaction. Drag, edit, create, comment, and delete operations update the interface immediately, persist to Supabase, and roll back visibly on failure.",
        "Realtime refresh. Postgres Changes refresh the owned workspace after task, comment, label, assignment, or activity events.",
    ]
    for item in architecture:
        add_list_item(doc, item, 78)

    doc.add_heading("Design decisions", level=2)
    design_items = [
        "Board-first hierarchy. The first viewport is the workflow itself, not generic dashboard chrome.",
        "Calm visual system. Warm ivory canvas, charcoal navigation, indigo focus color, and semantic slate/blue/amber/emerald workflow accents.",
        "Purposeful density. Metadata stays compact; task titles and status remain the strongest signals.",
        "Dedicated drag handle. Opening a task and moving it are distinct interactions, reducing accidental drags.",
        "Progressive detail. The right-side task drawer preserves board context on desktop and becomes a full-screen sheet on mobile.",
        "Honest state communication. Loading, empty, filtered-empty, reconnecting, mutation error, and demo states are deliberately visible.",
    ]
    for item in design_items:
        add_list_item(doc, item, 77)

    doc.add_heading("Implemented capabilities", level=1)
    add_feature_table(doc)

    doc.add_heading("Interaction details", level=2)
    for text in (
        "Drag and drop uses pointer activation distance, delayed touch activation, keyboard coordinates, a lifted overlay, destination highlighting, persistent position values, and rollback messaging.",
        "Search matches task title and description. Priority, assignee, and label filters combine with search, preserve the four-column spatial model, and report the visible result count.",
        "Due dates are overdue only for incomplete tasks before today; dates within two days are due soon; completed dates are visually subdued.",
        "Keyboard shortcuts focus search with slash or Command-K and open task creation with N when focus is not inside a form control.",
    ):
        add_list_item(doc, text, 77)

    doc.add_heading("Responsive behavior", level=2)
    doc.add_paragraph(
        "At wide widths all four columns fit beside the workspace rail. The rail collapses on tablet layouts. On mobile, columns become horizontal snap points at roughly 84% of the viewport, task details become full-screen, the search field remains reachable, and the compact workspace menu exposes team management without flattening the board into a generic list."
    )

    doc.add_heading("Supabase security model", level=1)
    add_callout(
        doc,
        "Security boundary",
        "Frontend visibility is not the authorization layer. Every table is protected in Postgres with RLS, grants, composite owner foreign keys, and ownership validation triggers. The service-role key is never required, stored, or committed.",
        DARK_BLUE,
    )
    security_items = [
        "Anonymous Supabase users receive the authenticated database role after sign-in; clients without a session keep the anon role and have no table privileges.",
        "Owned rows store user_id with a default of auth.uid(); policies require the stored value to equal the current session for SELECT, INSERT, UPDATE, and DELETE.",
        "Join tables store user_id and use composite foreign keys, so tasks cannot be linked to a member or label from another guest workspace.",
        "Activity history is append-only to browser clients. SECURITY DEFINER triggers generate authoritative entries for creation, updates, moves, assignments, labels, and comments.",
        "Realtime publication is registered idempotently. RLS SELECT policies remain the subscription boundary.",
        "Deleting an Auth user cascades through all owned product data; deleting a task cascades through its assignments, labels, comments, and history.",
    ]
    for item in security_items:
        add_list_item(doc, item, 77)

    doc.add_heading("Database entities", level=2)
    entities = [
        ("tasks", "Task content, workflow status, priority, due date, and numeric position."),
        ("team_members", "User-owned lightweight profiles for visual assignment."),
        ("task_assignees", "Many-to-many assignments with cross-owner protection."),
        ("comments", "Chronological task discussion with ownership validation."),
        ("labels", "Reusable custom names and colors."),
        ("task_labels", "Many-to-many task tagging with cross-owner protection."),
        ("activity_logs", "Trigger-authored, append-only task history."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    for idx, title in enumerate(("Table", "Responsibility")):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        set_cell_text(table.cell(0, idx), title, bold=True, color=DARK_BLUE, size=9)
    repeat_header_row(table.rows[0])
    for entity, purpose in entities:
        cells = table.add_row().cells
        set_cell_text(cells[0], entity, bold=True, size=9)
        set_cell_text(cells[1], purpose, size=9)
    set_table_geometry(table, [2700, 6660])

    doc.add_heading("Local setup", level=1)
    doc.add_paragraph("Prerequisites: Node.js 22.13 or newer and a free Supabase project.")
    setup_steps = [
        "Clone the public repository and install dependencies.",
        "Run the complete SQL in Appendix A (also available at supabase/schema.sql) in the Supabase SQL Editor.",
        "Enable Anonymous Sign-Ins under Authentication > Providers in the Supabase dashboard.",
        "Copy the project URL and public anon/publishable key into .env.local. Never add a service-role key.",
        "Start the development server and open http://localhost:3000.",
    ]
    for step in setup_steps:
        add_list_item(doc, step, 79)
    add_code_block(
        doc,
        [
            "git clone https://github.com/pratham-mavle/drift-task-board.git",
            "cd drift-task-board",
            "npm install",
            "cp .env.example .env.local",
            "npm run dev",
        ],
    )
    doc.add_heading("Environment values", level=2)
    add_code_block(
        doc,
        [
            "NEXT_PUBLIC_SUPABASE_URL=https://your-project-ref.supabase.co",
            "NEXT_PUBLIC_SUPABASE_ANON_KEY=your-public-anon-key",
        ],
    )
    add_callout(
        doc,
        "Deployment note",
        "The public Sites deployment is connected to a Supabase Nano Free Tier project in US East (Ohio). Hosting contains only the project URL and public publishable key. Anonymous Auth, RLS persistence, trigger-authored activity, and Realtime are active; no service-role or database credential is deployed.",
        GREEN,
    )

    doc.add_heading("Verification and delivery", level=2)
    checks = [
        "TypeScript compilation completed with no errors.",
        "ESLint completed with no errors or warnings.",
        "Production Vinext/Vite build completed successfully.",
        "Automated worker-rendering and security-contract tests passed (2 of 2).",
        "The complete migration was applied successfully to the hosted Supabase project with all seven tables, triggers, grants, RLS policies, and Realtime publication entries.",
        "A live two-session integration test passed anonymous sign-in, owner defaults, User A/User B isolation, cross-owner rejection, realtime task updates, comments, labels, assignments, status changes, activity logs, and cleanup.",
        "The Supabase-connected public Sites deployment returned HTTP 200 with the expected Drift metadata and application shell.",
        "Public GitHub repository contains source, setup guidance, environment example, and full schema; no secret environment file or service-role key is committed.",
    ]
    for check in checks:
        add_list_item(doc, check, 77)

    doc.add_heading("Tradeoffs and next steps", level=1)
    tradeoffs = [
        "Lightweight team profiles, not invited identities. This keeps anonymous onboarding simple. A multi-user version would add organizations, invitations, membership roles, and per-board permissions.",
        "One board per guest. A boards table plus board_id foreign keys would support multiple projects, templates, and archived boards.",
        "Client-side direct data access. It is appropriate because RLS is the security boundary. A server API or RPC becomes useful for complex atomic workflows, notifications, integrations, and rate limiting.",
        "Reordering rewrites affected positions. Fractional ranking or a database reorder RPC would reduce writes on very large columns.",
        "Plain-text comments. Mentions, rich text, attachments, notification delivery, and threaded replies are natural extensions.",
        "Automated QA covers compilation, linting, production rendering, schema behavior, and two independent live anonymous sessions. The next layer would drive drag, touch, focus trapping, and realtime reconnection in full browser tests.",
    ]
    for item in tradeoffs:
        add_list_item(doc, item, 77)

def add_sql_appendix(doc: Document):
    # Named override: the full SQL reference uses landscape Letter, 0.55-inch
    # margins, and Consolas 7pt so complete lines stay readable and uncut.
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    set_header_footer(section, "Appendix A  ·  Full Supabase schema", "Executable SQL  ·  RLS + realtime")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Appendix A — Full Supabase database schema")
    set_run_font(run, size=16, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("Source: supabase/schema.sql  |  Run once in the Supabase SQL Editor, then enable Anonymous Sign-Ins.")
    set_run_font(run, size=8.5, color=MUTED, italic=True)

    sql_text = (PROJECT_ROOT / "supabase" / "schema.sql").read_text(encoding="utf-8")
    for line in sql_text.splitlines():
        p = doc.add_paragraph(style="SQL Code")
        add_paragraph_shading(p, "F7F7F9")
        p_pr = p._p.get_or_add_pPr()
        widow = OxmlElement("w:widowControl")
        widow.set(qn("w:val"), "0")
        p_pr.append(widow)
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=7, color="24272F")


def build(output_path: Path):
    doc = Document()
    style_document(doc)
    props = doc.core_properties
    props.title = "Drift Kanban Task Manager — Project Delivery Report"
    props.subject = "Solution overview, setup, security model, features, tradeoffs, and full Supabase schema"
    props.author = "Pratham Mavle"
    props.keywords = "Kanban, React, TypeScript, Supabase, anonymous auth, RLS"
    props.comments = "Generated for the Drift task-manager project delivery."
    add_cover(doc)
    add_main_content(doc)
    add_sql_appendix(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: build_report.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
